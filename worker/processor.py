import logging
from collections import Counter
import re
import threading
from pathlib import Path

import pdfplumber
from docx import Document

from .ml_model import CareerModel
from .nlp_pipeline import extract_keywords
from .text_utils import build_text_stats, count_words, normalize_extracted_text

logger = logging.getLogger(__name__)


class WorkerProcessingError(RuntimeError):
    pass


class CVProcessor:
    SECTION_ALIASES = {
        "summary": ["summary", "profile", "objective", "about"],
        "experience": ["experience", "work experience", "employment", "professional experience"],
        "education": ["education", "studies", "academic"],
        "skills": ["skills", "technical skills", "core skills", "competencies"],
        "projects": ["projects", "portfolio", "personal projects"],
        "certifications": ["certifications", "certificates", "licenses"],
        "languages": ["languages", "language skills"],
    }

    ROLE_TERMS = [
        "developer",
        "engineer",
        "designer",
        "analyst",
        "scientist",
        "manager",
        "specialist",
        "consultant",
        "administrator",
        "architect",
        "intern",
    ]

    SKILL_TERMS = [
        "python",
        "django",
        "flask",
        "fastapi",
        "react",
        "javascript",
        "typescript",
        "node.js",
        "next.js",
        "html",
        "css",
        "tailwind",
        "postgresql",
        "mysql",
        "mongodb",
        "sql",
        "docker",
        "kubernetes",
        "rabbitmq",
        "redis",
        "aws",
        "azure",
        "gcp",
        "git",
        "github",
        "machine learning",
        "deep learning",
        "nlp",
        "pandas",
        "numpy",
        "scikit-learn",
        "tensorflow",
        "pytorch",
        "figma",
        "ux",
        "ui",
        "wireframes",
        "prototypes",
        "analytics",
        "excel",
        "power bi",
        "tableau",
        "testing",
        "api",
        "rest",
        "ci/cd",
    ]

    ACTION_VERBS = [
        "built",
        "created",
        "led",
        "improved",
        "reduced",
        "increased",
        "managed",
        "designed",
        "implemented",
        "deployed",
        "optimized",
        "analyzed",
        "automated",
        "delivered",
        "trained",
        "mentored",
        "coordinated",
        "launched",
        "maintained",
        "integrated",
    ]

    def __init__(self):
        self.career_model = CareerModel()

    def process(self, cv_id, file_path, user_id):
        aggregate = {
            "cv_id": cv_id,
            "user_id": user_id,
            "extracted_text": "",
            "keywords": [],
            "ml": {},
            "quality": {},
            "job_matches": [],
            "text_stats": {},
        }
        errors = []
        lock = threading.Lock()
        text_ready = threading.Event()
        barrier = threading.Barrier(4)

        def add_error(stage, exc):
            logger.exception("%s failed for CV %s", stage, cv_id)
            with lock:
                errors.append(f"{stage}: {exc}")

        def has_errors():
            with lock:
                return bool(errors)

        def wait_for_siblings():
            try:
                barrier.wait(timeout=120)
            except threading.BrokenBarrierError:
                logger.error("Processing barrier broke for CV %s", cv_id)

        def get_text():
            with lock:
                return aggregate["extracted_text"]

        def extract_thread():
            try:
                text, extraction_method = self.extract_text(file_path)
                if not text.strip():
                    raise WorkerProcessingError("No readable text was found in the uploaded file.")
                text_stats = build_text_stats(text, extraction_method)
                with lock:
                    aggregate["extracted_text"] = text
                    aggregate["text_stats"] = text_stats
            except Exception as exc:
                add_error("text extraction", exc)
            finally:
                text_ready.set()
                wait_for_siblings()

        def keyword_thread():
            text_ready.wait()
            try:
                if not has_errors():
                    keywords = extract_keywords(
                        get_text(),
                        self.career_model.job_corpus,
                        self.career_model.preprocessor,
                    )
                    with lock:
                        aggregate["keywords"] = keywords
            except Exception as exc:
                add_error("keyword extraction", exc)
            finally:
                wait_for_siblings()

        def scoring_thread():
            text_ready.wait()
            try:
                if not has_errors():
                    ml_result = self.career_model.score_text(get_text())
                    quality_result = self.career_model.score_quality(get_text())
                    with lock:
                        aggregate["ml"] = ml_result
                        aggregate["quality"] = quality_result
            except Exception as exc:
                add_error("ML scoring", exc)
            finally:
                wait_for_siblings()

        def matching_thread():
            text_ready.wait()
            try:
                if not has_errors():
                    matches = self.career_model.match_jobs(get_text())
                    with lock:
                        aggregate["job_matches"] = matches
            except Exception as exc:
                add_error("job matching", exc)
            finally:
                wait_for_siblings()

        threads = [
            threading.Thread(target=extract_thread, name=f"cv-{cv_id}-extract"),
            threading.Thread(target=keyword_thread, name=f"cv-{cv_id}-keywords"),
            threading.Thread(target=scoring_thread, name=f"cv-{cv_id}-ml-score"),
            threading.Thread(target=matching_thread, name=f"cv-{cv_id}-job-match"),
        ]

        for thread in threads:
            thread.start()
        for thread in threads:
            thread.join()

        with lock:
            if errors:
                raise WorkerProcessingError("; ".join(errors))
            ml = aggregate["ml"]
            quality = aggregate["quality"]
            matches = aggregate["job_matches"]
            keywords = aggregate["keywords"]
            extracted_text = aggregate["extracted_text"]
            text_stats = aggregate["text_stats"]
        insights = self.build_extended_insights(extracted_text, ml, quality, keywords, matches)

        return {
            "score": ml["score"],
            "predicted_category": ml["predicted_category"],
            "feedback": self.build_feedback(ml, keywords, matches),
            "cv_quality_score": quality["cv_quality_score"],
            "cv_quality_level": quality["cv_quality_level"],
            "cv_quality_feedback": quality["cv_quality_feedback"],
            "cv_quality_suggestions": quality["cv_quality_suggestions"],
            "cv_quality_breakdown": quality["cv_quality_breakdown"],
            "analysis_summary": insights["analysis_summary"],
            "personalization_profile": insights["personalization_profile"],
            "personalized_recommendations": insights["personalized_recommendations"],
            "strengths": insights["strengths"],
            "missing_keywords": insights["missing_keywords"],
            "career_path": insights["career_path"],
            "improvement_plan": insights["improvement_plan"],
            "rewrite_examples": insights["rewrite_examples"],
            "keywords": keywords,
            "job_matches": matches,
            "extracted_text": extracted_text,
            "text_stats": text_stats,
        }

    def extract_text(self, file_path):
        path = Path(file_path)
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            return self._extract_pdf(path)
        if suffix == ".docx":
            return self._extract_docx(path)
        raise WorkerProcessingError("Unsupported CV file type.")

    def _extract_pdf(self, path):
        pages = []
        methods = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text, method = self._extract_pdf_page(page)
                if text:
                    pages.append(text)
                    methods.append(method)
        method = self._most_common_method(methods, "pdf-empty")
        return normalize_extracted_text("\n\n".join(pages)), method

    def _extract_pdf_page(self, page):
        candidates = []

        plain_text = page.extract_text() or ""
        candidates.append(("pdf-text", plain_text))

        try:
            layout_text = page.extract_text(layout=True, x_tolerance=1, y_tolerance=3) or ""
            candidates.append(("pdf-layout", layout_text))
        except TypeError:
            logger.debug("pdfplumber layout extraction is unavailable for this page.")

        words_by_position = page.extract_words(x_tolerance=1, y_tolerance=3, keep_blank_chars=False)
        candidates.append(("pdf-words-position", self._pdf_words_to_lines(words_by_position)))

        words_by_flow = page.extract_words(use_text_flow=True, keep_blank_chars=False)
        candidates.append(("pdf-words-flow", " ".join(word.get("text", "") for word in words_by_flow)))

        return self._best_text_candidate(candidates)

    def _extract_docx(self, path):
        document = Document(path)
        blocks = []
        self._append_docx_paragraphs(blocks, document.paragraphs)
        self._append_docx_tables(blocks, document.tables)

        for section in document.sections:
            self._append_docx_paragraphs(blocks, section.header.paragraphs)
            self._append_docx_tables(blocks, section.header.tables)
            self._append_docx_paragraphs(blocks, section.footer.paragraphs)
            self._append_docx_tables(blocks, section.footer.tables)

        structured_text = "\n".join(blocks)
        xml_text = self._extract_docx_xml_text(document)
        text, method = self._best_text_candidate(
            [
                ("docx-structured", structured_text),
                ("docx-xml", xml_text),
            ]
        )
        return normalize_extracted_text(text), method

    def _append_docx_paragraphs(self, blocks, paragraphs):
        for paragraph in paragraphs:
            text = normalize_extracted_text(paragraph.text)
            if text:
                blocks.append(text)

    def _append_docx_tables(self, blocks, tables):
        for table in tables:
            for row in table.rows:
                cells = [normalize_extracted_text(cell.text) for cell in row.cells if cell.text.strip()]
                if cells:
                    blocks.append(" | ".join(cells))

    def _extract_docx_xml_text(self, document):
        elements = [document.element]
        for section in document.sections:
            elements.append(section.header._element)
            elements.append(section.footer._element)

        blocks = []
        for element in elements:
            for paragraph in element.iter():
                if paragraph.tag.rsplit("}", 1)[-1] != "p":
                    continue

                pieces = []
                for node in paragraph.iter():
                    tag = node.tag.rsplit("}", 1)[-1]
                    if tag in {"t", "instrText"} and node.text:
                        pieces.append(node.text)
                    elif tag == "tab":
                        pieces.append(" ")
                    elif tag == "br":
                        pieces.append("\n")

                text = normalize_extracted_text("".join(pieces))
                if text:
                    blocks.append(text)
        return "\n".join(blocks)

    def _pdf_words_to_lines(self, words, y_tolerance=3):
        if not words:
            return ""

        sorted_words = sorted(words, key=lambda word: (float(word.get("top", 0)), float(word.get("x0", 0))))
        lines = []
        current_line = []
        current_top = None

        for word in sorted_words:
            top = float(word.get("top", 0))
            if current_top is None or abs(top - current_top) <= y_tolerance:
                current_line.append(word)
                current_top = top if current_top is None else current_top
                continue

            lines.append(self._join_pdf_line(current_line))
            current_line = [word]
            current_top = top

        if current_line:
            lines.append(self._join_pdf_line(current_line))

        return "\n".join(line for line in lines if line)

    def _join_pdf_line(self, words):
        return " ".join(word.get("text", "") for word in sorted(words, key=lambda item: float(item.get("x0", 0))))

    def _best_text_candidate(self, candidates):
        normalized_candidates = [
            (method, normalize_extracted_text(text), count_words(text))
            for method, text in candidates
            if text and normalize_extracted_text(text)
        ]
        if not normalized_candidates:
            return "", "empty"

        method, text, _word_count = max(normalized_candidates, key=lambda item: (item[2], len(item[1])))
        return text, method

    def _most_common_method(self, methods, fallback):
        if not methods:
            return fallback
        return Counter(methods).most_common(1)[0][0]

    def build_feedback(self, ml, keywords, matches):
        top_keywords = ", ".join(item["term"] for item in keywords[:8]) or "no dominant keywords"
        best_match = matches[0] if matches else None
        match_text = (
            f"The closest role match is {best_match['title']} with {best_match['similarity']}% similarity."
            if best_match
            else "No close role match was found."
        )
        return (
            f"Score: {ml['score']} out of 100. This score is the model's confidence that the CV clearly fits "
            f"one of the trained career categories. Predicted category: {ml['predicted_category']} "
            f"with {int(round(ml['confidence'] * 100))}% category confidence. "
            f"The strongest TF-IDF keywords found in the CV are: {top_keywords}. {match_text} "
            "For a stronger score, make the target role explicit, repeat important role-specific tools naturally, "
            "and add measurable achievements that connect experience to the predicted category."
        )

    def build_extended_insights(self, text, ml, quality, keywords, matches):
        top_keywords = [item["term"] for item in keywords[:10]]
        best_match = matches[0] if matches else None
        profile = self.build_personalization_profile(text, ml, quality, keywords, matches)
        missing_keywords = self.find_missing_keywords(text, best_match, top_keywords)
        strengths = self.detect_strengths(quality, keywords, matches, profile)
        personalized_recommendations = self.build_personalized_recommendations(
            ml,
            quality,
            matches,
            missing_keywords,
            profile,
        )
        career_path = self.build_career_path(ml, quality, matches, missing_keywords, profile)
        improvement_plan = self.build_improvement_plan(quality, missing_keywords, best_match, profile)
        rewrite_examples = self.build_rewrite_examples(ml, missing_keywords, best_match, profile)

        best_match_text = (
            f" The closest specific job profile is {best_match['title']} at {best_match['similarity']}% similarity."
            if best_match
            else ""
        )
        skill_text = ", ".join(profile["detected_skills"][:5]) or "not enough explicit tools"
        evidence_text = (
            f"{profile['metric_count']} measurable result signals"
            if profile["metric_count"]
            else "no strong measurable result signals"
        )
        analysis_summary = (
            f"The analysis separates career direction from CV-writing quality. Career fit is {ml['score']}/100 for "
            f"{ml['predicted_category']}, while CV writing quality is {quality['cv_quality_score']}/100. "
            f"The strongest visible signals are {', '.join(top_keywords[:5]) or 'not yet clear'}, with explicit tools such as {skill_text} "
            f"and {evidence_text}.{best_match_text} The best next move is to strengthen the weakest writing-quality areas "
            "and add missing role vocabulary naturally."
        )

        return {
            "analysis_summary": analysis_summary,
            "personalization_profile": profile,
            "personalized_recommendations": personalized_recommendations,
            "strengths": strengths,
            "missing_keywords": missing_keywords,
            "career_path": career_path,
            "improvement_plan": improvement_plan,
            "rewrite_examples": rewrite_examples,
        }

    def build_personalization_profile(self, text, ml, quality, keywords, matches):
        lower_text = text.lower()
        detected_sections = self.detect_sections(lower_text)
        missing_sections = [name for name in self.SECTION_ALIASES if name not in detected_sections]
        detected_skills = self.detect_terms(text, self.SKILL_TERMS)
        detected_roles = self.detect_roles(text)
        action_verbs = self.detect_terms(text, self.ACTION_VERBS)
        metrics = self.extract_metrics(text)
        links = self.extract_links(text)
        top_keywords = [item["term"] for item in keywords[:8]]
        best_match = matches[0] if matches else None
        weak_quality_areas = sorted(quality["cv_quality_breakdown"], key=lambda item: item["score"])[:3]

        target_role = best_match["title"] if best_match else f"{ml['predicted_category']} role"
        target_terms = self.target_role_terms(best_match)
        missing_target_terms = [
            term
            for term in target_terms
            if term not in detected_skills and not self.contains_term(text, term)
        ][:10]

        return {
            "target_role": target_role,
            "predicted_category": ml["predicted_category"],
            "detected_roles": detected_roles[:8],
            "detected_skills": detected_skills[:14],
            "top_keywords": top_keywords,
            "detected_sections": detected_sections,
            "missing_sections": missing_sections[:5],
            "metric_count": len(metrics),
            "metrics": metrics[:8],
            "action_verbs": action_verbs[:10],
            "links": links[:5],
            "weak_quality_areas": weak_quality_areas,
            "missing_target_terms": missing_target_terms,
            "best_match_similarity": best_match["similarity"] if best_match else None,
        }

    def build_personalized_recommendations(self, ml, quality, matches, missing_keywords, profile):
        recommendations = []
        target_role = profile["target_role"]
        detected_skills = profile["detected_skills"]
        missing_terms = profile["missing_target_terms"] or [item["term"] for item in missing_keywords[:6]]
        weak_areas = profile["weak_quality_areas"]
        metric_count = profile["metric_count"]

        if detected_skills:
            recommendations.append(
                {
                    "priority": "High",
                    "title": f"Turn {', '.join(detected_skills[:3])} into proof for {target_role}",
                    "why": "The CV already names these skills, but the model rewards skills that are tied to responsibility and outcome.",
                    "how": f"Add one bullet for each important skill: what you built, where you used it, and what changed because of it.",
                    "example": f"Built a {target_role.lower()} project using {detected_skills[0]} to solve a specific problem, then report users, speed, accuracy, or delivery impact.",
                }
            )

        if missing_terms:
            recommendations.append(
                {
                    "priority": "High",
                    "title": f"Add missing language for {target_role}",
                    "why": f"The closest job profile expects signals such as {', '.join(missing_terms[:5])}, but they are not prominent in the CV.",
                    "how": "Only add terms you can honestly support. Connect every new keyword to a real project, course, responsibility, or tool choice.",
                    "example": f"Relevant project: used {', '.join(missing_terms[:3])} to deliver a concrete feature, analysis, workflow, or design decision.",
                }
            )

        if metric_count < 3:
            recommendations.append(
                {
                    "priority": "High",
                    "title": "Add more measurable outcomes",
                    "why": f"The CV currently shows {metric_count} measurable result signals. Recruiters can trust achievements faster when they include scale.",
                    "how": "Add percentages, user counts, response time, saved hours, project duration, team size, budget, or number of delivered features.",
                    "example": "Improved an internal workflow by reducing manual steps from X to Y, saving Z hours per week.",
                }
            )
        else:
            recommendations.append(
                {
                    "priority": "Medium",
                    "title": "Move the strongest metric into the top third",
                    "why": f"The CV already has measurable signals such as {', '.join(profile['metrics'][:3])}. They should be visible before the recruiter scans the details.",
                    "how": "Put the strongest metric in the summary or the first experience entry.",
                    "example": "Summary: role-focused candidate with experience delivering projects that improved X by Y%.",
                }
            )

        for area in weak_areas:
            recommendations.append(
                {
                    "priority": "High" if area["score"] < 55 else "Medium",
                    "title": f"Personalized fix for {area['name'].lower()}",
                    "why": f"This is one of the lower scoring writing areas: {area['evidence']}.",
                    "how": self.improvement_detail(area),
                    "example": self.personalized_area_example(area, target_role, detected_skills, missing_terms),
                }
            )

        if "projects" in profile["missing_sections"] and ml["predicted_category"] in {"Software Engineering", "Data Science", "Design"}:
            recommendations.append(
                {
                    "priority": "Medium",
                    "title": "Add a Projects section",
                    "why": f"For {target_role}, projects help prove ability when professional experience is short or broad.",
                    "how": "Add 2-3 compact projects with stack, goal, responsibility, result, and a link if available.",
                    "example": "Project name | stack | problem solved | measurable or visible outcome | repository/demo link.",
                }
            )

        if not profile["links"] and ml["predicted_category"] in {"Software Engineering", "Data Science", "Design"}:
            recommendations.append(
                {
                    "priority": "Medium",
                    "title": "Add proof links",
                    "why": "The model did not see portfolio, GitHub, LinkedIn, or project links. Proof links make claims easier to verify.",
                    "how": "Add one clean line near the header with GitHub, portfolio, LinkedIn, or a project demo.",
                    "example": "Links: GitHub | Portfolio | LinkedIn | live project demo.",
                }
            )

        recommendations.append(
            {
                "priority": "Medium",
                "title": f"Rewrite the summary for {target_role}",
                "why": "A role-specific summary helps both recruiters and the classifier understand the intended direction faster.",
                "how": "Use 2-3 lines: target role, strongest tools/domain, best proof, and the kind of work you want next.",
                "example": self.summary_example(target_role, detected_skills, profile["metrics"]),
            }
        )

        return recommendations[:10]

    def find_missing_keywords(self, text, best_match, existing_keywords):
        if not best_match:
            return []

        cv_terms = set(self.career_model.preprocessor.preprocess_tokens(text))
        existing_terms = set()
        for keyword in existing_keywords:
            existing_terms.update(keyword.split())

        job_terms = self.career_model.preprocessor.preprocess_tokens(best_match["description"])
        ranked_terms = Counter(job_terms).most_common()
        missing = []
        for term, frequency in ranked_terms:
            if term in cv_terms or term in existing_terms:
                continue
            if len(term) < 4:
                continue
            missing.append(
                {
                    "term": term,
                    "reason": f"Appears in the closest job profile ({best_match['title']}) but is not prominent in the CV.",
                    "priority": "high" if frequency > 1 else "medium",
                }
            )
            if len(missing) >= 8:
                break
        return missing

    def detect_sections(self, lower_text):
        detected = []
        for section, aliases in self.SECTION_ALIASES.items():
            for alias in aliases:
                pattern = rf"(^|\n)\s*{re.escape(alias)}\s*(:|\n|$)"
                if re.search(pattern, lower_text, flags=re.IGNORECASE):
                    detected.append(section)
                    break
        return detected

    def detect_terms(self, text, terms):
        detected = []
        for term in terms:
            if self.contains_term(text, term):
                detected.append(term)
        return detected

    def contains_term(self, text, term):
        if not text or not term:
            return False
        escaped = re.escape(term.lower()).replace(r"\ ", r"[\s/_-]+")
        pattern = rf"(?<![a-z0-9+#.-]){escaped}(?![a-z0-9+#.-])"
        return bool(re.search(pattern, text.lower()))

    def detect_roles(self, text):
        roles = []
        for line in text.splitlines():
            clean_line = normalize_extracted_text(line)
            if not clean_line or len(clean_line) > 90:
                continue
            lower_line = clean_line.lower()
            if any(role_term in lower_line for role_term in self.ROLE_TERMS):
                roles.append(clean_line)
        if roles:
            return list(dict.fromkeys(roles))

        detected_terms = [term for term in self.ROLE_TERMS if self.contains_term(text, term)]
        return detected_terms

    def extract_metrics(self, text):
        patterns = [
            r"\b\d+(?:[.,]\d+)?\s?%",
            r"\b\d+(?:[.,]\d+)?\s?(?:users|clients|customers|projects|features|tickets|reports|dashboards|models|pages|hours|days|weeks|months|years)\b",
            r"\b(?:reduced|increased|improved|saved|grew|optimized|delivered)\b[^.\n]{0,70}\b\d+(?:[.,]\d+)?\b",
            r"\$\s?\d+(?:[.,]\d+)?[kKmM]?",
            r"\b\d+(?:[.,]\d+)?x\b",
        ]
        metrics = []
        for pattern in patterns:
            for match in re.finditer(pattern, text, flags=re.IGNORECASE):
                metric = normalize_extracted_text(match.group())
                if metric and metric not in metrics:
                    metrics.append(metric)
                if len(metrics) >= 12:
                    return metrics
        return metrics

    def extract_links(self, text):
        links = []
        patterns = [
            r"https?://[^\s)>,]+",
            r"\b(?:github|linkedin|portfolio|behance|dribbble)\.com/[^\s)>,]+",
            r"\b[\w.+-]+@[\w-]+\.[\w.-]+\b",
        ]
        for pattern in patterns:
            for match in re.finditer(pattern, text, flags=re.IGNORECASE):
                value = match.group().strip(".,;")
                if value not in links:
                    links.append(value)
        return links

    def target_role_terms(self, best_match):
        if not best_match:
            return []
        job_terms = self.career_model.preprocessor.preprocess_tokens(best_match["description"])
        ranked_terms = [term for term, _count in Counter(job_terms).most_common()]
        return [term for term in ranked_terms if len(term) >= 4][:14]

    def personalized_area_example(self, area, target_role, detected_skills, missing_terms):
        main_skill = detected_skills[0] if detected_skills else (missing_terms[0] if missing_terms else "a relevant tool")
        if area["name"] == "Completeness":
            return f"Expanded a {target_role.lower()} responsibility by explaining the context, the {main_skill} work performed, and the final business or technical result."
        if area["name"] == "Structure":
            return "Use: Summary | Skills | Experience | Projects | Education, then keep each bullet to one action and one result."
        if area["name"] == "Measurable impact":
            return f"Improved a {target_role.lower()} workflow using {main_skill}, reducing time/cost/errors by X% or supporting Y users."
        if area["name"] == "Action language":
            return f"Built, optimized, automated, or delivered a {main_skill} solution for a clear {target_role.lower()} problem."
        if area["name"] == "Role clarity":
            return f"Headline: {target_role} focused on {main_skill}, delivery, and measurable outcomes."
        return f"Rewrite one generic sentence into a {target_role.lower()} achievement with tool, action, and result."

    def summary_example(self, target_role, detected_skills, metrics):
        skills = ", ".join(detected_skills[:4]) if detected_skills else "role-relevant tools"
        metric = metrics[0] if metrics else "measurable project outcomes"
        return f"{target_role} candidate with hands-on experience in {skills}. Strongest proof: {metric}. Looking to apply these skills in production work with clear ownership and impact."

    def detect_strengths(self, quality, keywords, matches, profile=None):
        strengths = []
        top_quality_areas = sorted(quality["cv_quality_breakdown"], key=lambda item: item["score"], reverse=True)[:2]
        for area in top_quality_areas:
            strengths.append(
                {
                    "title": area["name"],
                    "detail": f"{area['score']}/100 based on {area['evidence']}.",
                }
            )

        if keywords:
            strengths.append(
                {
                    "title": "Recognizable keyword signal",
                    "detail": f"Top extracted terms include {', '.join(item['term'] for item in keywords[:5])}.",
                }
            )

        if profile and profile["detected_skills"]:
            strengths.append(
                {
                    "title": "Concrete skill evidence",
                    "detail": f"The CV explicitly mentions {', '.join(profile['detected_skills'][:6])}. Tie these to outcomes for stronger recruiter value.",
                }
            )

        if profile and profile["metric_count"]:
            strengths.append(
                {
                    "title": "Some measurable proof",
                    "detail": f"Detected measurable signals such as {', '.join(profile['metrics'][:3])}.",
                }
            )

        if matches:
            strengths.append(
                {
                    "title": "Closest market direction",
                    "detail": f"{matches[0]['title']} is the closest role profile at {matches[0]['similarity']}% similarity.",
                }
            )
        return strengths[:7]

    def build_career_path(self, ml, quality, matches, missing_keywords, profile=None):
        target_role = matches[0]["title"] if matches else f"{ml['predicted_category']} role"
        second_role = matches[1]["title"] if len(matches) > 1 else "adjacent role"
        missing_terms = ", ".join(item["term"] for item in missing_keywords[:4]) or "the most important target-role tools"
        quality_focus = min(quality["cv_quality_breakdown"], key=lambda item: item["score"])
        detected_skills = ", ".join((profile or {}).get("detected_skills", [])[:4]) or "the strongest current skills"

        return [
            {
                "stage": "Now",
                "title": f"Position the CV toward {target_role}",
                "detail": f"Make the profile headline and first experience bullets clearly support {ml['predicted_category']} using {detected_skills}.",
            },
            {
                "stage": "Next 2 weeks",
                "title": "Close vocabulary gaps",
                "detail": f"Add honest, experience-backed references to {missing_terms}. Do not keyword-stuff; connect each term to a project or responsibility.",
            },
            {
                "stage": "Next project",
                "title": "Create proof for the target role",
                "detail": f"Add one portfolio or work project that demonstrates the responsibilities expected from a {target_role}.",
            },
            {
                "stage": "Interview positioning",
                "title": f"Prepare a bridge story for {second_role}",
                "detail": f"Use the CV to explain why your strongest experience transfers into {target_role} and nearby roles such as {second_role}.",
            },
            {
                "stage": "CV polish",
                "title": f"Improve {quality_focus['name'].lower()}",
                "detail": f"This is currently the weakest writing area: {quality_focus['evidence']}. Improving it should lift the writing-quality mark.",
            },
        ]

    def build_improvement_plan(self, quality, missing_keywords, best_match, profile=None):
        plan = []
        profile = profile or {}
        low_areas = sorted(quality["cv_quality_breakdown"], key=lambda item: item["score"])[:3]
        for area in low_areas:
            plan.append(
                {
                    "priority": "High" if area["score"] < 55 else "Medium",
                    "title": f"Improve {area['name'].lower()}",
                    "detail": self.improvement_detail(area),
                }
            )

        if missing_keywords:
            terms = ", ".join(item["term"] for item in missing_keywords[:5])
            role = best_match["title"] if best_match else "the target role"
            plan.append(
                {
                    "priority": "High",
                    "title": "Add missing role signals",
                    "detail": f"For {role}, the CV should naturally prove experience with: {terms}.",
                }
            )

        if profile.get("missing_sections"):
            plan.append(
                {
                    "priority": "Medium",
                    "title": "Add missing CV sections",
                    "detail": f"The parser did not clearly detect: {', '.join(profile['missing_sections'][:4])}. Add only the sections that are relevant and keep them compact.",
                }
            )

        plan.append(
            {
                "priority": "Medium",
                "title": "Move strongest evidence upward",
                "detail": "Put the strongest quantified achievements in the summary or first two experience bullets so they are visible immediately.",
            }
        )
        return plan[:7]

    def improvement_detail(self, area):
        if area["name"] == "Completeness":
            return "Expand short roles with scope, tools, project context, and outcome. Aim for enough detail that a recruiter can understand what you owned."
        if area["name"] == "Structure":
            return "Use consistent headings and bullet points. Keep each bullet to one idea: action, tool, result."
        if area["name"] == "Measurable impact":
            return "Add numbers such as percentages, time saved, users served, budget, team size, or frequency."
        if area["name"] == "Action language":
            return "Start bullets with verbs such as built, led, improved, automated, optimized, analyzed, or delivered."
        if area["name"] == "Role clarity":
            return "Add a headline or summary that names the target role and mirrors the most relevant skills."
        return "Add clearer evidence and reduce generic wording."

    def build_rewrite_examples(self, ml, missing_keywords, best_match, profile=None):
        role = best_match["title"] if best_match else f"{ml['predicted_category']} role"
        terms = [item["term"] for item in missing_keywords[:3]]
        term_text = ", ".join(terms) if terms else "the relevant tools"
        profile = profile or {}
        detected_skill = (profile.get("detected_skills") or ["a relevant tool"])[0]
        metric_hint = (profile.get("metrics") or ["a measurable result"])[0]
        return [
            {
                "before": "Responsible for different tasks and helping the team.",
                "after": f"Delivered {role.lower()} responsibilities by owning a defined project, using {term_text}, and reporting {metric_hint}.",
            },
            {
                "before": "Worked on applications and fixed issues.",
                "after": f"Improved reliability by identifying recurring issues, implementing fixes with {detected_skill}, and documenting before/after metrics.",
            },
            {
                "before": "Good communication and teamwork skills.",
                "after": "Coordinated with product, design, and technical stakeholders to clarify requirements, remove blockers, and deliver a documented project milestone.",
            },
        ]
